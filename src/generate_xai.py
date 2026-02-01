import os
import numpy as np
import matplotlib.pyplot as plt
import tensorflow as tf
from tensorflow.keras.preprocessing import image
from tensorflow.keras.models import load_model
from docx import Document
from docx.shared import Inches

# folders
data_dir = "data/combined"
output_dir = "outputs/xai"
os.makedirs(output_dir, exist_ok=True)

# your models
models = {
    "inception": "outputs/inception_best.h5",
    "xception": "outputs/xception_best.h5"
}

# one sample per class
classes = ["colon_cancer", "colon_normal", "polyp"]
samples = {}

for cls in classes:
    folder = os.path.join(data_dir, cls)
    files = [f for f in os.listdir(folder) if f.lower().endswith((".jpg", ".png", ".jpeg"))]
    samples[cls] = os.path.join(folder, files[0])

def gradcam(model, img_path, img_size=299):
    img = image.load_img(img_path, target_size=(img_size,img_size))
    x = image.img_to_array(img)
    x = np.expand_dims(x, axis=0)
    x = tf.keras.applications.inception_v3.preprocess_input(x)

    preds = model.predict(x)
    class_idx = np.argmax(preds[0])

    # pick last conv layer
    last_conv = None
    for layer in reversed(model.layers):
        if "conv" in layer.name:
            last_conv = layer.name
            break

    grad_model = tf.keras.models.Model(
        [model.inputs],
        [model.get_layer(last_conv).output, model.output]
    )

    with tf.GradientTape() as tape:
        conv_outputs, predictions = grad_model(x)
        loss = predictions[:, class_idx]

    grads = tape.gradient(loss, conv_outputs)
    pooled = tf.reduce_mean(grads, axis=(0,1,2))

    conv_outputs = conv_outputs[0]
    heatmap = conv_outputs @ pooled[..., tf.newaxis]
    heatmap = tf.squeeze(heatmap)
    heatmap = tf.maximum(heatmap, 0) / tf.math.reduce_max(heatmap)

    # resize & overlay
    img_orig = image.load_img(img_path)
    img_orig = image.img_to_array(img_orig)

    heatmap = np.uint8(255 * heatmap.numpy())
    heatmap = np.stack([heatmap]*3, axis=-1)
    heatmap = tf.image.resize(heatmap, (img_orig.shape[0], img_orig.shape[1])).numpy()
    overlay = (0.4*heatmap + img_orig).astype("uint8")

    return overlay

doc = Document()
doc.add_heading("XAI Grad-CAM Results", level=1)

for name, model_path in models.items():
    print(f"Loading model: {name}")
    model = load_model(model_path)
    doc.add_heading(name.upper(), level=2)

    for cls in classes:
        img_path = samples[cls]
        print(f"Processing {cls}")

        img = gradcam(model, img_path)
        out_path = os.path.join(output_dir, f"{name}_{cls}.jpg")
        plt.imsave(out_path, img)

        doc.add_heading(cls, level=3)
        doc.add_picture(out_path, width=Inches(4))

doc_path = "outputs/xai_results.docx"
doc.save(doc_path)
print(f"Done! Saved: {doc_path}")
